"""
Mock email provider.

Serves the messages defined in app.seed.mock_data and renders REAL attachment
bytes on demand (PDF / DOCX / PNG) so the in-app preview works exactly as it
will with Graph. No network, no credentials.
"""
from __future__ import annotations

import io

from app.seed import mock_data
from app.services.email_provider.base import (
    EmailProvider,
    ProviderAttachment,
    ProviderMessage,
)

_MONTHS = ["", "January", "February", "March", "April", "May", "June",
           "July", "August", "September", "October", "November", "December"]


# ---------- attachment renderers ----------
def _render_timesheet_pdf(case: dict) -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    title = "WEEKLY TIMESHEET" if case.get("period_label") else "MONTHLY TIMESHEET"
    pdf.cell(0, 10, title, ln=True)
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 8, f"Employee Name: {case.get('emp_name') or '(not printed)'}", ln=True)
    pdf.cell(0, 8, f"Employee ID: {case.get('emp_id') or '(not printed)'}", ln=True)
    pdf.cell(0, 8, f"Month: {_MONTHS[case['month']]} {case['year']}", ln=True)
    if case.get("period_label"):
        pdf.cell(0, 8, f"Period: {case['period_label']}", ln=True)
    pdf.ln(4)

    rows: list[tuple[str, str]] = []
    for d in case.get("annual", []):
        rows.append((d, "Annual Leave (AL)"))
    for d in case.get("remote", []):
        rows.append((d, "Work From Home (WFH)"))
    for d in case.get("sick", []):
        rows.append((d, "Sick Leave (SL)"))
    for d in case.get("unpaid", []):
        rows.append((d, "Unpaid Leave (LOP)"))
    for d in case.get("absent", []):
        rows.append((d, "Absent"))
    for d in case.get("public_holiday", []):
        rows.append((d, "Public Holiday (PH)"))
    rows.sort()

    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(60, 8, "Date", border=1)
    pdf.cell(80, 8, "Status", border=1, ln=True)
    pdf.set_font("Helvetica", "", 11)
    if not rows:
        pdf.cell(140, 8, "No leave recorded this month.", border=1, ln=True)
    for d, status in rows:
        pdf.cell(60, 8, d, border=1)
        pdf.cell(80, 8, status, border=1, ln=True)

    out = pdf.output(dest="S")
    data = bytes(out) if isinstance(out, (bytes, bytearray)) else out.encode("latin-1")
    if case.get("protected"):
        data = _encrypt_pdf(data)
    return data


def _encrypt_pdf(pdf_bytes: bytes, password: str = "secret123") -> bytes:
    """Password-protect a PDF so the pipeline's protection check can be demoed."""
    import fitz  # PyMuPDF

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        return doc.tobytes(
            encryption=fitz.PDF_ENCRYPT_AES_256,
            owner_pw=password, user_pw=password,
        )
    finally:
        doc.close()


def _render_timesheet_docx(case: dict) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Monthly Timesheet", level=1)
    doc.add_paragraph(f"Employee Name: {case.get('emp_name') or '(not printed)'}")
    doc.add_paragraph(f"Employee ID: {case.get('emp_id') or '(not printed)'}")
    doc.add_paragraph(f"Month: {_MONTHS[case['month']]} {case['year']}")
    if case.get("period_label"):
        doc.add_paragraph(f"Period: {case['period_label']}")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "Date", "Status"

    pairs: list[tuple[str, str]] = []
    for d in case.get("annual", []):
        pairs.append((d, "Annual Leave (AL)"))
    for d in case.get("sick", []):
        pairs.append((d, "Sick Leave (SL)"))
    for d in case.get("public_holiday", []):
        pairs.append((d, "Public Holiday (PH)"))
    for d, status in sorted(pairs):
        row = table.add_row().cells
        row[0].text, row[1].text = d, status

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _render_approval_png(detail: str, emp_name: str) -> bytes:
    from PIL import Image, ImageDraw

    w, h = 760, 240
    img = Image.new("RGB", (w, h), (245, 247, 250))
    d = ImageDraw.Draw(img)
    d.rectangle([12, 12, w - 12, h - 12], outline=(180, 188, 200), width=2)
    d.text((34, 38), "From: Sarah Khan (Account Manager)", fill=(40, 48, 60))
    d.text((34, 70), f"Subject: RE: Leave approval - {emp_name}", fill=(40, 48, 60))
    d.text((34, 120), "Confirmed. The leave dates below are", fill=(30, 30, 30))
    d.text((34, 150), detail, fill=(8, 110, 60))
    d.text((34, 196), "Regards, Sarah", fill=(90, 96, 105))
    bio = io.BytesIO()
    img.save(bio, format="PNG")
    return bio.getvalue()


def _case_filename(c: dict, ext: str) -> str:
    """Unique per attachment — weekly sheets for the same person/month must
    not collide (the period label keeps Week 1-2 and Week 3-4 apart)."""
    base = (c.get("emp_name") or "Unknown").replace(" ", "_").replace(".", "")
    period = f"_{c['period_label'].split(' ')[0].replace('-', 'to')}" if c.get("period_label") else ""
    return f"{base}_{_MONTHS[c['month']] or 'NoMonth'}_{c['year']}{period}_{c['slot']}.{ext}"


def _build_attachments(msg: dict) -> list[ProviderAttachment]:
    atts: list[ProviderAttachment] = []
    for c in msg["cases"]:
        is_docx = c["doc"] == "docx"
        ext = "docx" if is_docx else "pdf"
        ctype = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if is_docx else "application/pdf"
        )
        atts.append(ProviderAttachment(
            attachment_id=mock_data.attachment_id(msg["message_id"], c["slot"]),
            filename=_case_filename(c, ext),
            content_type=ctype,
            size=0,
            kind="timesheet",
        ))
    ap = msg.get("approval")
    if ap:
        atts.append(ProviderAttachment(
            attachment_id=mock_data.attachment_id(msg["message_id"], ap["slot"]),
            filename="manager_approval.png",
            content_type="image/png",
            size=0,
            kind="approval_screenshot",
        ))
    return atts


def _to_provider_message(msg: dict) -> ProviderMessage:
    return ProviderMessage(
        message_id=msg["message_id"],
        sender_name=msg["sender_name"],
        sender_email=msg["sender_email"],
        subject=msg["subject"],
        received_at=msg["received_at"],
        body_text=msg["body_text"],
        attachments=_build_attachments(msg),
    )


class MockEmailProvider(EmailProvider):
    async def list_messages(self, query: str | None = None) -> list[ProviderMessage]:
        msgs = [_to_provider_message(m) for m in mock_data.MESSAGES]
        if query:
            q = query.lower().strip()
            msgs = [
                m for m in msgs
                if q in (m.subject or "").lower()
                or q in (m.sender_name or "").lower()
                or q in (m.sender_email or "").lower()
                or q in (m.body_text or "").lower()
            ]
        return sorted(msgs, key=lambda m: m.received_at, reverse=True)

    async def get_message(self, message_id: str) -> ProviderMessage | None:
        raw = mock_data.message_by_id(message_id)
        return _to_provider_message(raw) if raw else None

    async def get_attachment_bytes(self, message_id: str, attachment_id: str) -> tuple[bytes, str, str]:
        msg = mock_data.message_by_id(message_id)
        if not msg:
            raise FileNotFoundError(message_id)

        # approval screenshot?
        ap = msg.get("approval")
        if ap and mock_data.attachment_id(message_id, ap["slot"]) == attachment_id:
            data = _render_approval_png(ap["detail"], msg["cases"][0]["emp_name"])
            return data, "manager_approval.png", "image/png"

        # otherwise a timesheet case
        case = mock_data.case_for_attachment(attachment_id)
        if not case:
            raise FileNotFoundError(attachment_id)
        if case["doc"] == "docx":
            data = _render_timesheet_docx(case)
            return data, _case_filename(case, "docx"), \
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        data = _render_timesheet_pdf(case)
        return data, _case_filename(case, "pdf"), "application/pdf"
