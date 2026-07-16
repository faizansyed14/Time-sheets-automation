"""
Converts uploads to JPEG images for the vision model, and extracts plain text
for grounding the model (and the deterministic leave/date checks).
"""
from __future__ import annotations

import calendar
import datetime as _dt
import io
import re
import subprocess
import tempfile
import zipfile
from html import unescape
from pathlib import Path

from app.core.config import settings

PDF_DPI = 300          # text-extraction default; the LLM render uses settings.pdf_render_dpi
PDF_MAX_PAGES = 10
JPEG_QUALITY = 90
IMAGE_MAX_SIDE = 2200
# Tall stitched PDF/DOCX screenshots — keep text readable (don't crush to 2200).
STITCH_MAX_WIDTH = 2200
STITCH_MAX_HEIGHT = 8000
# Outlook pastes attendance sheets into the HTML body as inline CID images.
_INLINE_TIMESHEET_MIN_BYTES = 20_000   # keep — real pasted grids
_INLINE_JUNK_MAX_BYTES = 12_000         # skip — signature logos/icons
_CID_REF_RE = re.compile(r"cid:([^\"'\s>)]+)", re.IGNORECASE)
_GENERIC_INLINE_NAME_RE = re.compile(
    r"^(image\d{2,3}\.(png|jpe?g|gif)|outlook-.+\.(png|jpe?g|gif|bmp)"
    r"|c2_signature_.+\.(png|jpe?g|gif))$", re.I)


# ---------------------------------------------------------------------------
# Shared date / attendance-grid scanning (used by the mock + vision engines)
# ---------------------------------------------------------------------------
_MONTH_NUM = {m.lower(): i for i, m in enumerate(calendar.month_name) if m}
_MONTH_NUM.update({m.lower(): i for i, m in enumerate(calendar.month_abbr) if m})

_ISO_RE = re.compile(r"\b(\d{4})-(\d{2})-(\d{2})\b")
_DMY_RE = re.compile(r"\b(\d{1,2})[ \-/]+([A-Za-z]{3,9})[ \-/,]+(\d{4})\b")
_MDY_RE = re.compile(r"\b([A-Za-z]{3,9})[ \-/]+(\d{1,2})[ \-/,]+(\d{4})\b")
# A clock time like "8:30 AM" / "17:00" → marks a day as worked/present.
_TIME_RE = re.compile(r"\b\d{1,2}:\d{2}\s*(?:[AaPp]\.?[Mm]\.?)?\b")
_WEEKEND_WORDS = ("weekend", "wekend", "week end", "rest day", "off day", "weekly off")


def _emit_iso(y: int, mo: int | None, d: int) -> str | None:
    try:
        return _dt.date(y, mo, d).isoformat() if mo else None
    except (ValueError, TypeError):
        return None


def find_dates_in_text(text: str) -> list[tuple[int, int, str]]:
    """Return (start, end, iso_date) for every ISO / 'DD Mon YYYY' / 'Mon DD, YYYY'
    date in the text, de-duplicated and ordered by position."""
    out: list[tuple[int, int, str]] = []
    for m in _ISO_RE.finditer(text):
        iso = _emit_iso(int(m[1]), int(m[2]), int(m[3]))
        if iso:
            out.append((m.start(), m.end(), iso))
    for m in _DMY_RE.finditer(text):
        iso = _emit_iso(int(m[3]), _MONTH_NUM.get(m[2].lower()), int(m[1]))
        if iso:
            out.append((m.start(), m.end(), iso))
    for m in _MDY_RE.finditer(text):
        iso = _emit_iso(int(m[3]), _MONTH_NUM.get(m[1].lower()), int(m[2]))
        if iso:
            out.append((m.start(), m.end(), iso))
    out.sort()
    deduped: list[tuple[int, int, str]] = []
    last_end = -1
    for s, e, iso in out:
        if s >= last_end:
            deduped.append((s, e, iso))
            last_end = e
    return deduped


def scan_attendance_grid(text: str) -> tuple[set[str], set[str]]:
    """For a daily-grid timesheet, return (present_days, weekend_days) as ISO
    date sets. 'present' = a clock time / hours appear next to the date;
    'weekend' = a weekend marker appears. Used to detect working days that have
    NEITHER hours NOR a leave entry (i.e. unaccounted days)."""
    present: set[str] = set()
    weekend: set[str] = set()
    dates = find_dates_in_text(text or "")
    for i, (_s, e, iso) in enumerate(dates):
        seg_end = dates[i + 1][0] if i + 1 < len(dates) else min(len(text), e + 60)
        segment = text[e:seg_end]
        low = segment.lower()
        if _TIME_RE.search(segment):
            present.add(iso)
        elif any(w in low for w in _WEEKEND_WORDS):
            weekend.add(iso)
    return present, weekend


def _to_jpeg_bytes(img, quality: int = JPEG_QUALITY) -> bytes:
    buf = io.BytesIO()
    img.convert("RGB").save(buf, format="JPEG", quality=quality, optimize=True)
    return buf.getvalue()


def detect_file_type(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if data.startswith(b"%PDF"):
        return "pdf"
    if data.startswith(b"\x89PNG\r\n\x1a\n") or data.startswith(b"\xff\xd8\xff"):
        return "image"
    if data[:2] == b"PK" and zipfile.is_zipfile(io.BytesIO(data)):
        if name.endswith(".docx"):
            return "docx"
        if name.endswith(".xlsx"):
            return "xlsx"
    if name.endswith(".pdf"):
        return "pdf"
    if name.endswith(".docx"):
        return "docx"
    if name.endswith(".xlsx"):
        return "xlsx"
    if name.endswith(".eml"):
        return "eml"
    if name.endswith((".png", ".jpg", ".jpeg")):
        return "image"
    return "unknown"


def image_to_images(image_bytes: bytes, max_side: int = IMAGE_MAX_SIDE) -> list[bytes]:
    from PIL import Image

    img = Image.open(io.BytesIO(image_bytes)).convert("RGB")
    w, h = img.size
    scale = min(1.0, max_side / max(w, h))
    if scale < 1.0:
        img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
    return [_to_jpeg_bytes(img)]


def _soffice_to_pdf(in_path: Path, out_dir: Path) -> bytes | None:
    """Run LibreOffice headless to convert any supported file to PDF.

    Uses a per-call private user profile so concurrent ingests don't collide on
    the default profile lock (a common cause of silent failures)."""
    profile = out_dir / "_lo_profile"
    cmd = ["soffice", "--headless", "--nologo", "--nofirststartwizard", "--norestore",
           f"-env:UserInstallation=file://{profile}",
           "--convert-to", "pdf", "--outdir", str(out_dir), str(in_path)]
    try:
        subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL,
                       stderr=subprocess.DEVNULL, timeout=120)
    except Exception:
        return None
    pdfs = [p for p in out_dir.glob("*.pdf")]
    return pdfs[0].read_bytes() if pdfs else None


def _office_to_pdf_bytes(file_bytes: bytes, ext: str) -> bytes | None:
    """Convert docx/xlsx -> PDF via LibreOffice headless, if soffice is installed."""
    if not file_bytes or ext not in {"docx", "xlsx"}:
        return None
    with tempfile.TemporaryDirectory() as td:
        td_path = Path(td)
        in_path = td_path / f"input.{ext}"
        out_dir = td_path / "out"
        out_dir.mkdir(parents=True, exist_ok=True)
        in_path.write_bytes(file_bytes)
        return _soffice_to_pdf(in_path, out_dir)


def _weasyprint_html_to_pdf(html: str) -> bytes | None:
    """Render HTML+CSS to PDF with WeasyPrint — a true visual render that keeps
    table borders, background colours and layout (a real 'screenshot' of the
    email). Pure-Python, no browser needed. Returns None if unavailable."""
    try:
        from weasyprint import HTML
    except Exception:
        return None

    # Don't hit the network: inline (cid:) logos and remote/tracking images are
    # returned as empty so rendering is fast, offline and safe. data: URIs are
    # handled natively by WeasyPrint.
    def _fetch(url: str):
        if url.startswith("data:"):
            from weasyprint.urls import default_url_fetcher
            return default_url_fetcher(url)
        return {"string": b"", "mime_type": "image/png"}

    wrapper = (
        "<style>@page{size:A4;margin:10mm;} "
        "body{font-family:Arial,Helvetica,'DejaVu Sans',sans-serif;} "
        "table{border-collapse:collapse;} img{max-width:100%;} "
        # keep table rows whole so nothing is cut across a page break
        "tr,td,th{break-inside:avoid;page-break-inside:avoid;}</style>"
    )
    try:
        return HTML(string=wrapper + html, url_fetcher=_fetch).write_pdf()
    except Exception:
        return None


def _html_to_pdf_bytes(html: str) -> bytes | None:
    """Render an HTML document to PDF, preferring a high-fidelity engine.
      1) WeasyPrint  — keeps colours, table lines, layout (pure-Python)
      2) LibreOffice — headless Writer/Web (works where soffice is installed)
    Returns None if neither is available so callers fall back to a text image."""
    if not (html or "").strip():
        return None
    pdf = _weasyprint_html_to_pdf(html)
    if pdf:
        return pdf
    with tempfile.TemporaryDirectory() as td:
        td_path = Path(td)
        in_path = td_path / "email.html"
        out_dir = td_path / "out"
        out_dir.mkdir(parents=True, exist_ok=True)
        in_path.write_text(html, encoding="utf-8")
        return _soffice_to_pdf(in_path, out_dir)


# Legible TrueType font for text-render fallbacks (much clearer than PIL's
# default bitmap font, which is a major reason rendered text was unreadable).
_FONT_PATHS = (
    "DejaVuSans.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
)
_FONT_CACHE: dict[int, object] = {}


def _load_font(size: int):
    if size in _FONT_CACHE:
        return _FONT_CACHE[size]
    from PIL import ImageFont
    font = None
    for p in _FONT_PATHS:
        try:
            font = ImageFont.truetype(p, size)
            break
        except Exception:
            continue
    if font is None:
        font = ImageFont.load_default()
    _FONT_CACHE[size] = font
    return font


def docx_to_images(docx_bytes: bytes) -> list[bytes]:
    from PIL import Image, ImageDraw

    pdf = _office_to_pdf_bytes(docx_bytes, "docx")
    if pdf:
        return [pdf_to_single_image(pdf)]
    # fallback: render extracted text
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    text = "\n".join(lines[:400]) or "(empty DOCX)"
    img = Image.new("RGB", (1800, 2400), (255, 255, 255))
    d = ImageDraw.Draw(img)
    y = 20
    for ln in text.splitlines():
        d.text((20, y), ln[:240], fill=(0, 0, 0))
        y += 20
    return [_to_jpeg_bytes(img)]


def xlsx_to_images(xlsx_bytes: bytes) -> list[bytes]:
    from PIL import Image, ImageDraw

    pdf = _office_to_pdf_bytes(xlsx_bytes, "xlsx")
    if pdf:
        return [pdf_to_single_image(pdf)]
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(xlsx_bytes), data_only=True)
    ws = wb.active
    lines = []
    for r in range(1, min(ws.max_row or 1, 400) + 1):
        vals = [("" if ws.cell(r, c + 1).value is None else str(ws.cell(r, c + 1).value))
                for c in range(min(ws.max_column or 1, 40))]
        if any(vals):
            lines.append(" | ".join(vals))
    text = "\n".join(lines) or "(empty XLSX)"
    img = Image.new("RGB", (2400, 3000), (255, 255, 255))
    d = ImageDraw.Draw(img)
    y = 20
    for ln in text.splitlines():
        d.text((20, y), ln[:340], fill=(0, 0, 0))
        y += 18
    return [_to_jpeg_bytes(img)]


_TIMESHEET_MARKERS = (
    "emp no", "employee id", "employee no", "emp id", "subject: timesheet", "timesheet -",
)


def _parse_eml(eml_bytes: bytes):
    from email import policy
    from email.parser import BytesParser

    return BytesParser(policy=policy.default).parsebytes(eml_bytes)


def _html_to_text(html: str) -> str:
    """Flatten HTML tables into pipe-separated rows for the vision / text steps."""
    h = html or ""
    h = re.sub(r"(?i)</tr\s*>", "\n", h)
    h = re.sub(r"(?i)</t[dh]\s*>", " | ", h)
    h = re.sub(r"(?i)<br\s*/?>", "\n", h)
    h = re.sub(r"<[^>]+>", "", h)
    text = unescape(h)
    lines: list[str] = []
    for ln in text.splitlines():
        ln = re.sub(r"\s+", " ", ln).strip()
        if ln:
            lines.append(ln)
    return "\n".join(lines)


def _score_timesheet_text(text: str) -> int:
    low = (text or "").lower()
    score = 0
    for marker in _TIMESHEET_MARKERS:
        if marker in low:
            score += 50
    for kw in ("sick leave", "annual leave", "public holiday", "weekend", "work from home", "wfh"):
        score += low.count(kw) * 5
    score += min(len(text) // 200, 40)
    return score


def _extract_eml_body_text(eml_bytes: bytes) -> str:
    """Best-effort body text from an .eml (plain + HTML parts, reply chains included)."""
    try:
        msg = _parse_eml(eml_bytes)
    except Exception:
        return ""

    chunks: list[str] = []
    for part in msg.walk():
        if part.get_content_maintype() == "multipart":
            continue
        ct = part.get_content_type()
        try:
            if ct == "text/plain":
                chunks.append(part.get_content())
            elif ct == "text/html":
                chunks.append(_html_to_text(part.get_content()))
        except Exception:
            continue

    if not chunks:
        return ""

    return max(chunks, key=_score_timesheet_text).strip()


def _focus_timesheet_text(text: str) -> str:
    """Drop reply headers / signatures; keep the forwarded timesheet block."""
    lines = text.splitlines()
    start = 0
    for i, ln in enumerate(lines):
        low = ln.lower().strip()
        if any(m in low for m in _TIMESHEET_MARKERS):
            start = i
            break
    focused = "\n".join(lines[start:]).strip()
    return focused or text.strip()


def _part_file_type(part, payload: bytes) -> str | None:
    """Resolve file type from filename, content-type, or magic bytes."""
    filename = part.get_filename() or ""
    ftype = detect_file_type(filename, payload)
    if ftype != "unknown":
        return ftype
    ct = (part.get_content_type() or "").lower()
    if ct == "application/pdf" or payload.startswith(b"%PDF"):
        return "pdf"
    if ct in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return "docx"
    if ct in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ):
        return "xlsx"
    if ct.startswith("image/"):
        return "image"
    return None


def _score_eml_attachment(filename: str, payload: bytes, ftype: str) -> int:
    fn = (filename or "").lower()
    score = 0
    if any(k in fn for k in ("timesheet", "time sheet", "time-sheet")):
        score += 120
    if ftype == "pdf":
        score += 60
    elif ftype in ("docx", "xlsx"):
        score += 50
    elif ftype == "image":
        score += 20
    score += min(len(payload) // 8000, 40)
    return score


def _eml_collect_attachments(eml_bytes: bytes) -> list[tuple[str, bytes, str]]:
    """Return (filename, payload, file_type) for real attachments inside .eml.

    Excluded:
    - Small CID-referenced inline images (logos, icons embedded in HTML).
    - Other tiny inline images without an attachment disposition.

    Kept:
    - PDF / Office / nested .eml parts — even when they carry a Content-Id
      (Outlook and Graph often tag file attachments with CIDs too).
    - Parts with Content-Disposition: attachment.
    - Large inline images (e.g. approval screenshots)."""
    try:
        msg = _parse_eml(eml_bytes)
    except Exception:
        return []

    found: list[tuple[str, bytes, str]] = []
    for part in msg.walk():
        maintype = part.get_content_maintype()
        if maintype == "multipart":
            continue

        # A forwarded email (message/rfc822) carries the real timesheet inside
        # itself ("email inside the email"). Python's walk() descends into it
        # automatically when it's MIME-nested, but when the nested message is
        # base64-encoded as a single leaf it is NOT descended — so pull its raw
        # bytes and recurse to dig out the attached PDF/sheet.
        if maintype == "message":
            nested = _nested_message_bytes(part)
            if nested:
                found.extend(_eml_collect_attachments(nested))
            continue

        payload = part.get_payload(decode=True)
        if not payload:
            continue

        filename = part.get_filename() or ""
        disposition = (part.get_content_disposition() or "").lower()
        cid = (part.get("Content-Id") or "").strip()
        ftype = _part_file_type(part, payload)
        if not ftype:
            continue

        # Documents are always real attachments — Graph/Outlook assign Content-Id
        # to PDFs and Office files, not only HTML-embedded images.
        if ftype in ("pdf", "docx", "xlsx", "eml"):
            found.append((filename, payload, ftype))
            continue

        if disposition == "attachment":
            found.append((filename, payload, ftype))
            continue

        if ftype == "image":
            fn = filename.lower()
            n = len(payload)
            # Tiny inline logos/icons — never a timesheet.
            if n < _INLINE_JUNK_MAX_BYTES and disposition != "attachment":
                continue
            # Pasted attendance grids in the body (Outlook inline CID) — keep even
            # when the filename is generic (image001.png).
            if (
                disposition == "attachment"
                or n >= _INLINE_TIMESHEET_MIN_BYTES
                or any(k in fn for k in (
                    "screenshot", "approval", "attendance", "timesheet",
                    "smarttime", "report", "capture",
                ))
            ):
                found.append((filename, payload, ftype))

    # De-duplicate by CONTENT, not name — Outlook/Graph commonly attach the
    # exact same image twice: once inline (cid-referenced, generic name like
    # "image005.png") and once as a plain attachment (its real name, e.g.
    # "ATT00008.png"). Hashing the bytes catches this even though the two
    # copies carry different filenames; when both names collide, keep the
    # non-generic one so the sheet isn't shown to the model/vault as
    # "body_timesheet.png" when a real name was available.
    import hashlib

    seen: dict[bytes, tuple[str, bytes, str]] = {}
    order: list[bytes] = []
    for fn, payload, ftype in found:
        key = hashlib.sha256(payload).digest()
        if key not in seen:
            seen[key] = (fn, payload, ftype)
            order.append(key)
            continue
        existing_fn, _existing_payload, _existing_ft = seen[key]
        if _GENERIC_INLINE_NAME_RE.match(existing_fn or "") and not _GENERIC_INLINE_NAME_RE.match(fn or ""):
            seen[key] = (fn, payload, ftype)
    deduped = [seen[k] for k in order]
    return deduped


def _nested_message_bytes(part) -> bytes | None:
    """Raw .eml bytes of a message/rfc822 part, whether MIME-nested or a
    base64-encoded leaf. Returns None if nothing usable is found."""
    try:
        payload = part.get_payload(decode=True)
        if payload and payload.lstrip()[:1] not in (b"", None):
            return payload
    except Exception:
        pass
    try:
        inner = part.get_payload()
        if isinstance(inner, list) and inner:
            return inner[0].as_bytes()
        if hasattr(inner, "as_bytes"):
            return inner.as_bytes()
        if isinstance(inner, str) and inner.strip():
            return inner.encode("utf-8", "replace")
    except Exception:
        pass
    return None


def _text_to_page_images(
    text: str,
    *,
    width: int = 1654,      # ~A4 at 200 DPI
    height: int = 2339,
    font_size: int = 26,
    margin: int = 60,
) -> list[bytes]:
    """Render plain text to clean, legible A4 page images using a TrueType font
    (with word wrapping). Used as the fallback when LibreOffice isn't available
    to render the email HTML."""
    from PIL import Image, ImageDraw

    font = _load_font(font_size)
    line_h = int(font_size * 1.4)
    raw_lines = [ln.rstrip() for ln in (text or "").splitlines()] or ["(empty email)"]

    # word-wrap to the page width
    probe = ImageDraw.Draw(Image.new("RGB", (10, 10)))
    max_w = width - 2 * margin

    def _wrap(line: str) -> list[str]:
        if not line:
            return [""]
        words, out, cur = line.split(" "), [], ""
        for w in words:
            trial = f"{cur} {w}".strip()
            if probe.textlength(trial, font=font) <= max_w:
                cur = trial
            else:
                if cur:
                    out.append(cur)
                cur = w
        out.append(cur)
        return out

    wrapped: list[str] = []
    for ln in raw_lines:
        wrapped.extend(_wrap(ln))

    lines_per_page = max(1, (height - 2 * margin) // line_h)
    images: list[bytes] = []
    for off in range(0, len(wrapped), lines_per_page):
        img = Image.new("RGB", (width, height), (255, 255, 255))
        d = ImageDraw.Draw(img)
        y = margin
        for ln in wrapped[off : off + lines_per_page]:
            d.text((margin, y), ln, fill=(20, 20, 20), font=font)
            y += line_h
        images.append(_to_jpeg_bytes(img))
    return images or [_to_jpeg_bytes(Image.new("RGB", (width, height), (255, 255, 255)))]


def _likely_pasted_timesheet_image(payload: bytes) -> bool:
    """Keep Outlook-pasted attendance grids; drop logos / marketing banners.

    Thresholds match attachment junk policy: tiny CID parts are signature
    chrome; very wide images are email footer banners (brand strips).
    """
    if not payload or len(payload) < _INLINE_TIMESHEET_MIN_BYTES:
        return False
    try:
        from PIL import Image

        im = Image.open(io.BytesIO(payload))
        w, h = im.size
        if max(w, h) < 280:
            return False
        # UAE / corporate banner strips are wide and short.
        if h > 0 and w >= h * 2:
            return False
        return True
    except Exception:
        return len(payload) >= 40_000


_IMG_TAG_CID_RE = re.compile(
    r"<img\b[^>]*\bsrc\s*=\s*[\"']cid:[^\"']+[\"'][^>]*/?>",
    re.IGNORECASE,
)


def _resolve_eml_html_cids(eml_bytes: bytes, html_body: str) -> str:
    """Replace cid: image refs in HTML with data: URIs from the same .eml.

    Only pasted timesheet-sized images are inlined. Signature logos and wide
    marketing banners stay out of the JPEG so OpenAI never sees brand chrome.
    Without inlining real grids, body→PDF renders show empty boxes.
    """
    import base64

    if not html_body or "cid:" not in html_body.lower():
        return html_body
    try:
        msg = _parse_eml(eml_bytes)
    except Exception:
        return html_body
    cid_map: dict[str, str] = {}
    for part in msg.walk():
        if part.get_content_maintype() != "image":
            continue
        cid = (part.get("Content-Id") or "").strip().strip("<>")
        if not cid:
            continue
        payload = part.get_payload(decode=True)
        if not payload or not _likely_pasted_timesheet_image(payload):
            continue
        ctype = (part.get_content_type() or "image/png").split(";")[0]
        uri = f"data:{ctype};base64,{base64.b64encode(payload).decode()}"
        cid_map[cid.lower()] = uri
        cid_map[cid.split("@")[0].lower()] = uri
        fn = (part.get_filename() or "").strip().lower()
        if fn:
            cid_map[fn] = uri

    def _cid_uri(ref: str) -> str | None:
        return (cid_map.get(ref.lower())
                or cid_map.get(ref.split("@")[0].lower()))

    def _sub_ref(m: re.Match) -> str:
        return _cid_uri(m.group(1)) or ""

    # Drop whole <img cid:...> tags that are logos/banners (not in cid_map).
    def _sub_img(m: re.Match) -> str:
        ref_m = _CID_REF_RE.search(m.group(0))
        if not ref_m:
            return ""
        return m.group(0) if _cid_uri(ref_m.group(1)) else ""

    out = _IMG_TAG_CID_RE.sub(_sub_img, html_body)
    return _CID_REF_RE.sub(_sub_ref, out)


_SHEET_SIG_LABEL_HTML_RE = re.compile(
    r"(?is)\b(?:EMPLOYEE|MANAGER)\s+SIGNATURE\b",
)


def _strip_html_after_sheet_signatures(html: str) -> str:
    """Cut corporate contact cards / logo blocks after sheet signature labels.

    Many DHRE sheets end with EMPLOYEE/MANAGER SIGNATURE then repeat brand
    banners + phones — that chrome must not become vision pixels.
    """
    if not html:
        return html
    matches = list(_SHEET_SIG_LABEL_HTML_RE.finditer(html))
    if not matches:
        return html
    end = matches[-1].end()
    tail = html[end:]
    tail_text = re.sub(r"<[^>]+>", "", tail).strip()
    if len(tail_text) < 20 and "cid:" not in tail.lower() and "<img" not in tail.lower():
        return html
    return html[:end] + "<!-- signature-redacted -->"


def _eml_subject_and_body_html(eml_bytes: bytes) -> tuple[str, str | None, str]:
    """Return (subject, html_body_or_None, plain_body) for an .eml."""
    msg = _parse_eml(eml_bytes)
    subject = msg.get("subject", "") or ""
    html_body: str | None = None
    plain_body = ""
    for part in msg.walk():
        if part.get_content_maintype() == "multipart":
            continue
        ct = part.get_content_type()
        try:
            if ct == "text/html" and html_body is None:
                html_body = part.get_content()
            elif ct == "text/plain" and not plain_body:
                plain_body = part.get_content()
        except Exception:
            continue
    return subject, html_body, plain_body


def _trim_whitespace(img, bg=(255, 255, 255), pad: int = 16):
    """Crop the surrounding white border off a rendered page image."""
    from PIL import Image, ImageChops

    bgimg = Image.new("RGB", img.size, bg)
    diff = ImageChops.difference(img.convert("RGB"), bgimg)
    bbox = diff.getbbox()
    if not bbox:
        return img
    l, t, r, b = bbox
    l, t = max(0, l - pad), max(0, t - pad)
    r, b = min(img.width, r + pad), min(img.height, b + pad)
    return img.crop((l, t, r, b))


def pdf_to_single_image(
    pdf_bytes: bytes,
    dpi: int | None = None,
    gap: int = 12,
    max_pages: int = PDF_MAX_PAGES,
) -> bytes:
    """Render a (possibly multi-page) PDF into ONE continuous tall image, with
    each page's whitespace trimmed and the pages stacked vertically. Because
    table rows are kept whole (CSS break-inside:avoid), nothing is cut mid-row,
    so the timesheet reads as a single uninterrupted screenshot for the LLM."""
    import fitz
    from PIL import Image

    if dpi is None:
        dpi = int(getattr(settings, "pdf_render_dpi", 150) or 150)
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    zoom = dpi / 72.0
    mat = fitz.Matrix(zoom, zoom)
    pages: list = []
    for i in range(min(doc.page_count, max_pages)):
        pix = doc.load_page(i).get_pixmap(matrix=mat, alpha=False)
        im = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        pages.append(_trim_whitespace(im))
    doc.close()
    if not pages:
        raise ValueError("empty PDF")
    if len(pages) == 1:
        canvas = pages[0]
    else:
        width = max(p.width for p in pages)
        height = sum(p.height for p in pages) + gap * (len(pages) - 1)
        canvas = Image.new("RGB", (width, height), (255, 255, 255))
        y = 0
        for p in pages:
            canvas.paste(p, (0, y))
            y += p.height + gap
    scale = min(STITCH_MAX_WIDTH / canvas.width, STITCH_MAX_HEIGHT / canvas.height, 1.0)
    if scale < 1.0:
        canvas = canvas.resize(
            (max(1, int(canvas.width * scale)), max(1, int(canvas.height * scale))),
            Image.LANCZOS,
        )
    return _to_jpeg_bytes(canvas)


def eml_body_to_images(eml_bytes: bytes) -> list[bytes]:
    """Render the EMAIL ITSELF (Subject + body) to ONE continuous tall image —
    the inline timesheet table renders crisply (WeasyPrint → LibreOffice) with
    full colours/borders, stitched into a single screenshot so nothing is split
    across a page break. Saved as evidence AND sent to the vision model. Falls
    back to a clean TrueType text render if no HTML engine is available."""
    from app.core.pii import scrub_email_for_llm, scrub_text

    try:
        subject, html_body, plain_body = _eml_subject_and_body_html(eml_bytes)
    except Exception:
        _, scrubbed = scrub_email_for_llm("", _extract_eml_body_text(eml_bytes))
        return [_stitch_text_images(scrubbed)]

    # Scrub BEFORE rendering — once addresses/phones/secrets are pixels in the
    # JPEG the vision model reads them like any other text. Signature footers
    # are cut so "Password:" lines under Thanks never leave the server.
    # CID logos / wide brand banners are not inlined (see _resolve_eml_html_cids).
    subject, plain_body = scrub_email_for_llm(subject, plain_body)
    if html_body:
        from app.core.pii import strip_html_quoted_reply_thread

        _, focused_plain = scrub_email_for_llm(subject, _extract_eml_body_text(eml_bytes))
        html_body = scrub_text(html_body)
        html_body = strip_html_quoted_reply_thread(html_body)
        html_body = _strip_html_after_sheet_signatures(html_body)
        html_body = _resolve_eml_html_cids(eml_bytes, html_body)
        # Prefer focused plain when HTML is reply-chain chrome (no real grid).
        # Do NOT treat the word "timesheet" in RE: subjects as an inline sheet —
        # that kept full quoted threads in the body JPEG.
        html_l = (html_body or "").lower()
        has_inline_grid = any(
            m in html_l
            for m in (
                "attendance sheet",
                "emp no",
                "empno",
                "hours worked",
                "daily total",
            )
        ) or (
            "<table" in html_l
            and any(m in html_l for m in ("public holiday", "leave", "regular", "in</", "out</"))
        )
        has_pasted_sheet_img = "data:image/" in html_l
        plain_trimmed = (
            "[signature-redacted]" in (focused_plain or "")
            or "[quoted-thread-redacted]" in (focused_plain or "")
        )
        if focused_plain and plain_trimmed and not has_inline_grid and not has_pasted_sheet_img:
            plain_body = focused_plain
            html_body = None
        else:
            plain_body = scrub_text(plain_body)
    header = f"<p style='font-family:sans-serif'><b>Subject:</b> {unescape_safe(subject)}</p><hr/>"
    if html_body:
        doc = (
            "<html><head><meta charset='utf-8'>"
            "<style>body{font-family:Arial,Helvetica,sans-serif;font-size:11pt;} "
            "table{border-collapse:collapse;} "
            "tr,td,th{break-inside:avoid;page-break-inside:avoid;}</style>"
            f"</head><body>{header}{html_body}</body></html>"
        )
    else:
        from html import escape as _esc
        focused = _focus_timesheet_text(plain_body)
        doc = f"<html><body>{header}<pre style='font-family:monospace;font-size:11pt'>{_esc(focused)}</pre></body></html>"

    pdf = _html_to_pdf_bytes(doc)
    if pdf:
        try:
            return [pdf_to_single_image(pdf)]
        except Exception:
            pass
    # Fallback: legible text render (no HTML engine available)
    _, body_fb = scrub_email_for_llm(subject, _extract_eml_body_text(eml_bytes))
    text = f"Subject: {subject}\n\n" + _focus_timesheet_text(body_fb)
    return [_stitch_text_images(text)]


def email_body_to_images(subject: str | None, body_text: str | None) -> list[bytes]:
    """Render a Graph/inbox plain-text email (subject + body) to JPEG — same
    pipeline path as inline EML timesheets. PII is scrubbed before rendering
    so addresses/phones/secrets never become pixels the vision model can read."""
    from html import escape as _esc

    from app.core.pii import scrub_email_for_llm

    subject, body = scrub_email_for_llm(subject, body_text)
    focused = _focus_timesheet_text(body)
    header = (
        f"<p style='font-family:sans-serif'><b>Subject:</b> "
        f"{_esc(subject or '')}</p><hr/>"
    )
    doc = (
        "<html><head><meta charset='utf-8'>"
        "<style>body{font-family:Arial,Helvetica,sans-serif;font-size:11pt;} "
        "pre{white-space:pre-wrap;word-wrap:break-word;font-family:monospace;font-size:10pt;} "
        "table{border-collapse:collapse;} "
        "tr,td,th{break-inside:avoid;page-break-inside:avoid;}</style>"
        f"</head><body>{header}<pre>{_esc(focused)}</pre></body></html>"
    )
    pdf = _html_to_pdf_bytes(doc)
    if pdf:
        try:
            return [pdf_to_single_image(pdf)]
        except Exception:
            pass
    text = f"Subject: {subject or ''}\n\n{focused}"
    return [_stitch_text_images(text)]


def _stitch_text_images(text: str) -> bytes:
    """Stitch the multi-page TrueType text render into one tall image."""
    from PIL import Image

    pages = [Image.open(io.BytesIO(b)).convert("RGB") for b in _text_to_page_images(text)]
    if len(pages) == 1:
        return _to_jpeg_bytes(_trim_whitespace(pages[0]))
    pages = [_trim_whitespace(p) for p in pages]
    width = max(p.width for p in pages)
    height = sum(p.height for p in pages) + 12 * (len(pages) - 1)
    canvas = Image.new("RGB", (width, height), (255, 255, 255))
    y = 0
    for p in pages:
        canvas.paste(p, (0, y))
        y += p.height + 12
    return _to_jpeg_bytes(canvas)


def unescape_safe(s: str) -> str:
    return unescape(s or "")


def eml_best_attachment(eml_bytes: bytes) -> tuple[str, bytes, str] | None:
    """Pick the single most likely timesheet file when an .eml has several attachments."""
    attachments = _eml_collect_attachments(eml_bytes)
    if not attachments:
        return None
    return max(attachments, key=lambda t: _score_eml_attachment(t[0], t[1], t[2]))


def eml_all_attachments(eml_bytes: bytes) -> list[tuple[str, bytes, str]]:
    """Every real (non-decorative) document attachment inside an .eml, as
    (save_name, payload, file_type). Used to file EACH attached sheet separately
    in the vault alongside the original .eml — not just the single best one.

    Save names are de-duplicated so two attachments that share a filename (or
    have none) never overwrite each other on disk."""
    raw = _eml_collect_attachments(eml_bytes)
    out: list[tuple[str, bytes, str]] = []
    used: set[str] = set()
    for idx, (fname, payload, ftype) in enumerate(raw, 1):
        name = eml_attachment_save_name(fname, ftype, payload)
        if name in used:
            stem, dot, ext = name.rpartition(".")
            name = f"{stem}_{idx}{dot}{ext}" if dot else f"{name}_{idx}"
        used.add(name)
        out.append((name, payload, ftype))
    return out


def eml_attachment_save_name(filename: str, ftype: str, payload: bytes | None = None) -> str:
    name = (filename or "").strip()
    n = len(payload or b"")
    if name and not _GENERIC_INLINE_NAME_RE.match(name):
        return name
    if ftype == "image" and n >= _INLINE_TIMESHEET_MIN_BYTES:
        return "body_timesheet.png"
    if name:
        return name
    ext = {"pdf": ".pdf", "docx": ".docx", "xlsx": ".xlsx", "image": ".png"}.get(ftype, "")
    return f"extracted_timesheet{ext}"


def _eml_attachment_images(eml_bytes: bytes) -> list[bytes]:
    """PDF / Office / large inline images attached to the email (not logos)."""
    best = eml_best_attachment(eml_bytes)
    if not best:
        return []
    _name, payload, ftype = best
    try:
        return to_images(ftype, payload)
    except Exception:
        return []


def eml_to_images(eml_bytes: bytes) -> list[bytes]:
    # 1) a real attached PDF / Office sheet is the most faithful source
    attached = _eml_attachment_images(eml_bytes)
    if attached:
        return attached
    # 2) otherwise render the EMAIL ITSELF (Subject + body) — the inline
    #    timesheet table renders crisply via LibreOffice, with a clean
    #    TrueType text fallback if LibreOffice isn't available.
    return eml_body_to_images(eml_bytes)


def to_images(file_type: str, data: bytes) -> list[bytes]:
    """One JPEG per document for vision. PDFs/Office docs are stitched into a
    single tall screenshot (same idea as the email-body render)."""
    if file_type == "pdf":
        return [pdf_to_single_image(data)]
    if file_type == "docx":
        return docx_to_images(data)
    if file_type == "xlsx":
        return xlsx_to_images(data)
    if file_type == "eml":
        return eml_to_images(data)
    if file_type == "image":
        return image_to_images(data)
    raise ValueError(f"Unsupported file type: {file_type}")


def extract_document_text(file_type: str, data: bytes) -> str:
    """Plain text for the optional text cross-validation step."""
    try:
        if file_type == "pdf":
            import fitz
            doc = fitz.open(stream=data, filetype="pdf")
            parts = []
            for i in range(min(doc.page_count, PDF_MAX_PAGES)):
                parts.append(doc.load_page(i).get_text())
            doc.close()
            return "\n".join(parts).strip()
        if file_type == "docx":
            from docx import Document
            doc = Document(io.BytesIO(data))
            lines = [p.text for p in doc.paragraphs if p.text.strip()]
            for t in doc.tables:
                for row in t.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if any(cells):
                        lines.append(" | ".join(cells))
            return "\n".join(lines).strip()
        if file_type == "xlsx":
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(data), data_only=True)
            ws = wb.active

            def _fill_rgb(cell) -> str | None:
                """Solid fill colour as RRGGBB, or None for no/plain fill.
                Colour often carries the MEANING on timesheets (legend-coded
                leave), which plain text loses — so annotate it."""
                try:
                    f = cell.fill
                    if f is None or f.patternType != "solid":
                        return None
                    rgb = getattr(f.fgColor, "rgb", None)
                    if not isinstance(rgb, str) or len(rgb) < 6:
                        return None
                    rgb = rgb[-6:].upper()
                    return None if rgb in ("000000", "FFFFFF") else rgb
                except Exception:
                    return None

            lines = []
            for r in range(1, min(ws.max_row or 1, 500) + 1):
                vals = []
                row_fill: str | None = None
                for c in range(min(ws.max_column or 1, 50)):
                    cell = ws.cell(r, c + 1)
                    v = "" if cell.value is None else str(cell.value)
                    fill = _fill_rgb(cell)
                    if fill and not v:
                        row_fill = row_fill or fill
                    elif fill:
                        v = f"{v} [fill={fill}]"
                    vals.append(v)
                if any(vals):
                    line = " | ".join(vals)
                    if row_fill:
                        line += f"  [row fill={row_fill}]"
                    lines.append(line)
            # Merged ranges: the label lives only in the top-left cell — spell
            # out the span so "Eid Al Adha" across 5 rows reads as 5 days.
            try:
                merged = []
                for rng in list(ws.merged_cells.ranges)[:80]:
                    tl = ws.cell(rng.min_row, rng.min_col).value
                    if tl is not None and str(tl).strip() and rng.max_row > rng.min_row:
                        merged.append(f'MERGED {rng.coord}: "{str(tl).strip()[:60]}" '
                                      f"spans rows {rng.min_row}-{rng.max_row} "
                                      "(applies to every row in the span)")
                if merged:
                    lines.append("--- MERGED CELLS ---")
                    lines.extend(merged)
            except Exception:
                pass
            return "\n".join(lines).strip()
        if file_type == "eml":
            best = eml_best_attachment(data)
            if best:
                _name, payload, ftype = best
                doc_text = extract_document_text(ftype, payload)
                if doc_text.strip():
                    return doc_text.strip()
            return _focus_timesheet_text(_extract_eml_body_text(data))
    except Exception:
        return ""
    return ""
